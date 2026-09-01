from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from PIL import GifImagePlugin, Image
from pypdf import PdfWriter

from ai_fde.modules.evidence.parser import (
    MAX_ARCHIVE_EXPANDED_BYTES,
    MAX_ARCHIVE_FILES,
    MAX_CSV_COLUMNS,
    MAX_DOCUMENT_CHARACTERS,
    MAX_IMAGE_PIXELS,
    MAX_PDF_PAGES,
    EvidenceParseError,
    UnsupportedEvidenceTypeError,
    parse_evidence,
    parse_text_evidence,
    validate_evidence_upload_metadata,
)


def test_text_parser_retains_exact_source_offsets() -> None:
    source = b"Title\n\nInvoices over $50,000 require CFO approval.\n\nLast paragraph.\n"

    segments = parse_text_evidence(source, "text/markdown", "policy.md")

    assert [segment.content for segment in segments] == [
        "Title",
        "Invoices over $50,000 require CFO approval.",
        "Last paragraph.",
    ]
    decoded = source.decode()
    for segment in segments:
        assert decoded[segment.start_offset : segment.end_offset] == segment.content
        assert segment.locator == {
            "kind": "text_offset",
            "start": segment.start_offset,
            "end": segment.end_offset,
        }
        assert segment.parser_name == "utf8-paragraph-parser"


def test_csv_parser_produces_addressable_rows() -> None:
    segments = parse_evidence(
        b"Process,Owner,System\nInvoice approval,CFO,NetSuite\n",
        "text/csv",
        "workflow.csv",
    )

    assert [segment.content for segment in segments] == [
        "Process | Owner | System",
        "Invoice approval | CFO | NetSuite",
    ]
    assert segments[1].locator == {"kind": "csv_row", "row": 2}


def test_email_parser_preserves_headers_and_plain_text_body() -> None:
    source = (
        b"From: controller@example.test\r\n"
        b"To: ap@example.test\r\n"
        b"Subject: Approval exception\r\n"
        b"Content-Type: text/plain; charset=utf-8\r\n\r\n"
        b"Strategic vendors may be approved by the Controller.\r\n"
    )

    segments = parse_evidence(source, "message/rfc822", "approval.eml")

    assert segments[0].locator == {"kind": "email_headers"}
    assert "Subject: Approval exception" in segments[0].content
    assert segments[1].locator == {
        "kind": "email_body",
        "part": 0,
        "media_type": "text/plain",
    }
    assert "Strategic vendors" in segments[1].content


def test_docx_parser_preserves_paragraph_and_table_row_locators() -> None:
    document = Document()
    document.add_paragraph("Invoices over $50,000 require CFO approval.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "System"
    table.cell(0, 1).text = "NetSuite"
    buffer = io.BytesIO()
    document.save(buffer)

    segments = parse_evidence(
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "policy.docx",
    )

    assert segments[0].locator == {"kind": "docx_paragraph", "paragraph": 1}
    assert segments[1].locator == {"kind": "docx_table_row", "table": 1, "row": 1}
    assert segments[1].content == "System | NetSuite"


def test_image_parser_validates_content_and_anchors_the_whole_image() -> None:
    buffer = io.BytesIO()
    Image.new("RGB", (24, 12), color="white").save(buffer, format="PNG")

    [segment] = parse_evidence(buffer.getvalue(), "image/png", "diagram.png")

    assert segment.modality == "image"
    assert segment.locator == {"kind": "whole_image", "width": 24, "height": 12}
    assert "diagram.png" in segment.content


def test_parser_rejects_type_mismatch_malformed_content_and_non_utf8_text() -> None:
    with pytest.raises(UnsupportedEvidenceTypeError, match="does not match"):
        parse_evidence(b"not a pdf", "text/plain", "policy.pdf")

    with pytest.raises(EvidenceParseError, match="PDF evidence is malformed"):
        parse_evidence(b"%PDF malformed", "application/pdf", "policy.pdf")

    with pytest.raises(UnsupportedEvidenceTypeError, match="valid UTF-8"):
        parse_text_evidence(b"\xff\xfe", "text/plain", "invalid.txt")

    with pytest.raises(UnsupportedEvidenceTypeError, match="Supported evidence types"):
        parse_evidence(b"payload", "application/zip", "archive.zip")


def test_upload_metadata_validation_normalizes_extension_and_content_type() -> None:
    assert validate_evidence_upload_metadata("text/markdown; charset=utf-8", "Policy.MD") == "md"
    assert validate_evidence_upload_metadata("APPLICATION/PDF", "report.PDF") == "pdf"

    with pytest.raises(UnsupportedEvidenceTypeError, match="Supported evidence types"):
        validate_evidence_upload_metadata("text/plain", "notes")

    with pytest.raises(UnsupportedEvidenceTypeError, match="does not match"):
        validate_evidence_upload_metadata("image/png", "policy.txt")


def test_parser_rejects_evidence_with_no_readable_content() -> None:
    with pytest.raises(EvidenceParseError, match="did not contain any readable content"):
        parse_evidence(b"   \n\n  \n", "text/plain", "empty.txt")


def test_parser_rejects_text_exceeding_the_document_character_limit() -> None:
    oversized = b"x" * (MAX_DOCUMENT_CHARACTERS + 1)

    with pytest.raises(EvidenceParseError, match="exceeds the safe text limit"):
        parse_evidence(oversized, "text/plain", "huge.txt")


def test_csv_parser_rejects_too_many_columns() -> None:
    row = ",".join(["cell"] * (MAX_CSV_COLUMNS + 1)).encode()

    with pytest.raises(EvidenceParseError, match="exceeds the safe column limit"):
        parse_evidence(row, "text/csv", "wide.csv")


def test_csv_parser_rejects_malformed_quoting() -> None:
    with pytest.raises(EvidenceParseError, match="CSV evidence is malformed"):
        parse_evidence(b'Process,Owner\n"unterminated,CFO\x00\n', "text/csv", "broken.csv")


def test_csv_parser_skips_structurally_empty_rows() -> None:
    segments = parse_evidence(
        b"Process,Owner\n,\nInvoice approval,CFO\n", "text/csv", "process.csv"
    )

    assert [segment.locator["row"] for segment in segments] == [1, 3]


def test_pdf_parser_rejects_encrypted_documents() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.encrypt("evidence-password")
    buffer = io.BytesIO()
    writer.write(buffer)

    with pytest.raises(EvidenceParseError, match="Encrypted PDF evidence is not supported"):
        parse_evidence(buffer.getvalue(), "application/pdf", "secret.pdf")


def test_pdf_parser_rejects_documents_over_the_page_limit() -> None:
    writer = PdfWriter()
    for _ in range(MAX_PDF_PAGES + 1):
        writer.add_blank_page(width=72, height=72)
    buffer = io.BytesIO()
    writer.write(buffer)

    with pytest.raises(EvidenceParseError, match="exceeds the safe page limit"):
        parse_evidence(buffer.getvalue(), "application/pdf", "long.pdf")


def _zip_bytes(entries: dict[str, bytes]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, payload in entries.items():
            archive.writestr(name, payload)
    return buffer.getvalue()


DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_docx_parser_rejects_an_archive_that_expands_past_the_safe_limit() -> None:
    """A highly compressible member is rejected on its declared size, before extraction."""
    bomb = _zip_bytes({"word/document.xml": b"\0" * (MAX_ARCHIVE_EXPANDED_BYTES + 1)})
    assert len(bomb) < 100_000

    with pytest.raises(EvidenceParseError, match="exceeds the safe expansion limit"):
        parse_evidence(bomb, DOCX_CONTENT_TYPE, "bomb.docx")


def test_docx_parser_rejects_an_archive_containing_a_traversal_path() -> None:
    archive = _zip_bytes({"../../etc/passwd": b"payload"})

    with pytest.raises(EvidenceParseError, match="contains an unsafe path"):
        parse_evidence(archive, DOCX_CONTENT_TYPE, "traversal.docx")


def test_docx_parser_rejects_an_archive_with_too_many_members() -> None:
    archive = _zip_bytes({f"part-{index}.xml": b"x" for index in range(MAX_ARCHIVE_FILES + 1)})

    with pytest.raises(EvidenceParseError, match="exceeds the safe file limit"):
        parse_evidence(archive, DOCX_CONTENT_TYPE, "many.docx")


def test_docx_parser_rejects_content_that_is_not_a_zip_container() -> None:
    with pytest.raises(EvidenceParseError, match="DOCX evidence is malformed"):
        parse_evidence(b"plain text, not a container", DOCX_CONTENT_TYPE, "fake.docx")


def test_image_parser_rejects_images_over_the_safe_pixel_limit() -> None:
    buffer = io.BytesIO()
    Image.new("L", (6000, 5000)).save(buffer, format="PNG")
    assert MAX_IMAGE_PIXELS < 6000 * 5000

    with pytest.raises(EvidenceParseError, match="exceeds the safe pixel limit"):
        parse_evidence(buffer.getvalue(), "image/png", "huge.png")


def test_image_parser_rejects_truncated_image_bytes() -> None:
    buffer = io.BytesIO()
    Image.new("RGB", (32, 32), color="white").save(buffer, format="PNG")
    truncated = buffer.getvalue()[:40]

    with pytest.raises(EvidenceParseError):
        parse_evidence(truncated, "image/png", "truncated.png")


def test_text_only_entrypoint_rejects_binary_evidence_formats() -> None:
    with pytest.raises(UnsupportedEvidenceTypeError, match="must be UTF-8 text or Markdown"):
        parse_text_evidence(b"%PDF-1.4", "application/pdf", "report.pdf")
def test_image_parser_does_not_invoke_decoders_outside_the_declared_format() -> None:
    """A .png upload must not reach Pillow's other format plugins.

    The extension allowlist alone runs *after* `Image.open()`, so without an
    explicit `formats` restriction every registered decoder parses attacker
    bytes before the mismatch is detected.
    """
    buffer = io.BytesIO()
    Image.new("P", (4, 4)).save(buffer, format="GIF")
    disguised_gif = buffer.getvalue()
    assert disguised_gif[:6] in (b"GIF87a", b"GIF89a")

    entered: list[str] = []
    original_open = GifImagePlugin.GifImageFile._open

    def record_open(self: GifImagePlugin.GifImageFile) -> None:
        entered.append("GIF")
        original_open(self)

    with pytest.MonkeyPatch.context() as patch:
        patch.setattr(GifImagePlugin.GifImageFile, "_open", record_open)
        with pytest.raises(EvidenceParseError, match="does not match its file extension"):
            parse_evidence(disguised_gif, "image/png", "diagram.png")

    assert entered == []
