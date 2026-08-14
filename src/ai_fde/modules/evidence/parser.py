from __future__ import annotations

import csv
import io
import re
import zipfile
from dataclasses import dataclass
from email import policy
from email.message import Message
from email.parser import BytesParser
from html.parser import HTMLParser
from pathlib import Path
from typing import Literal

from docx import Document
from PIL import Image, UnidentifiedImageError
from pypdf import PdfReader
from pypdf.errors import PdfReadError


class UnsupportedEvidenceTypeError(ValueError):
    pass


class EvidenceParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedSegment:
    ordinal: int
    content: str
    start_offset: int
    end_offset: int
    locator: dict[str, int | str]
    parser_name: str
    parser_version: str = "1.0.0"
    modality: Literal["text", "image"] = "text"


MAX_PDF_PAGES = 100
MAX_CSV_ROWS = 10_000
MAX_CSV_COLUMNS = 200
MAX_DOCUMENT_CHARACTERS = 2_000_000
MAX_ARCHIVE_FILES = 1_000
MAX_ARCHIVE_EXPANDED_BYTES = 20 * 1024 * 1024
MAX_IMAGE_PIXELS = 25_000_000

SUPPORTED_EVIDENCE_EXTENSIONS = frozenset(
    {"txt", "md", "csv", "eml", "pdf", "docx", "png", "jpg", "jpeg"}
)

_CONTENT_TYPES: dict[str, frozenset[str]] = {
    "txt": frozenset({"text/plain", "application/octet-stream"}),
    "md": frozenset({"text/markdown", "text/plain", "application/octet-stream"}),
    "csv": frozenset({"text/csv", "application/csv", "application/octet-stream"}),
    "eml": frozenset({"message/rfc822", "application/octet-stream"}),
    "pdf": frozenset({"application/pdf", "application/octet-stream"}),
    "docx": frozenset(
        {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/octet-stream",
        }
    ),
    "png": frozenset({"image/png", "application/octet-stream"}),
    "jpg": frozenset({"image/jpeg", "application/octet-stream"}),
    "jpeg": frozenset({"image/jpeg", "application/octet-stream"}),
}


def validate_evidence_upload_metadata(content_type: str, file_name: str) -> str:
    suffix = Path(file_name).suffix.casefold().lstrip(".")
    normalized_type = content_type.casefold().split(";", maxsplit=1)[0].strip()
    if suffix not in SUPPORTED_EVIDENCE_EXTENSIONS:
        supported = ", ".join(f".{item}" for item in sorted(SUPPORTED_EVIDENCE_EXTENSIONS))
        raise UnsupportedEvidenceTypeError(f"Supported evidence types are: {supported}.")
    if normalized_type not in _CONTENT_TYPES[suffix]:
        raise UnsupportedEvidenceTypeError(
            "The uploaded content type does not match the evidence file extension."
        )
    return suffix


def parse_evidence(content: bytes, content_type: str, file_name: str) -> list[ParsedSegment]:
    suffix = validate_evidence_upload_metadata(content_type, file_name)
    try:
        if suffix in {"txt", "md"}:
            segments = _parse_text(content)
        elif suffix == "csv":
            segments = _parse_csv(content)
        elif suffix == "eml":
            segments = _parse_email(content)
        elif suffix == "pdf":
            segments = _parse_pdf(content)
        elif suffix == "docx":
            segments = _parse_docx(content)
        else:
            segments = _parse_image(content, suffix, file_name)
    except UnsupportedEvidenceTypeError:
        raise
    except EvidenceParseError:
        raise
    except Exception as exc:  # noqa: BLE001 - parser boundary returns a content-free failure
        raise EvidenceParseError("Evidence could not be parsed safely.") from exc
    if not segments:
        raise EvidenceParseError("Evidence did not contain any readable content.")
    total_characters = sum(len(segment.content) for segment in segments)
    if total_characters > MAX_DOCUMENT_CHARACTERS:
        raise EvidenceParseError("Parsed evidence exceeds the safe text limit.")
    return segments


def parse_text_evidence(content: bytes, content_type: str, file_name: str) -> list[ParsedSegment]:
    """Compatibility entrypoint retained for callers that explicitly require text evidence."""
    suffix = validate_evidence_upload_metadata(content_type, file_name)
    if suffix not in {"txt", "md"}:
        raise UnsupportedEvidenceTypeError("Evidence must be UTF-8 text or Markdown.")
    return _parse_text(content)


def _decode_utf8(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise UnsupportedEvidenceTypeError("Text evidence must be valid UTF-8.") from exc


def _parse_text(content: bytes) -> list[ParsedSegment]:
    text = _decode_utf8(content)
    segments: list[ParsedSegment] = []
    for match in re.finditer(
        r"\S(?:.*?\S)?(?=(?:\r?\n)[ \t]*(?:\r?\n)|\s*\Z)",
        text,
        flags=re.DOTALL,
    ):
        segment_text = match.group(0).strip()
        if not segment_text:
            continue
        start = text.find(segment_text, match.start(), match.end())
        end = start + len(segment_text)
        segments.append(
            ParsedSegment(
                ordinal=len(segments),
                content=segment_text,
                start_offset=start,
                end_offset=end,
                locator={"kind": "text_offset", "start": start, "end": end},
                parser_name="utf8-paragraph-parser",
            )
        )
    return segments


def _parse_csv(content: bytes) -> list[ParsedSegment]:
    text = _decode_utf8(content)
    try:
        rows = csv.reader(io.StringIO(text, newline=""), strict=True)
        segments: list[ParsedSegment] = []
        for row_number, row in enumerate(rows, start=1):
            if row_number > MAX_CSV_ROWS:
                raise EvidenceParseError("CSV evidence exceeds the safe row limit.")
            if len(row) > MAX_CSV_COLUMNS:
                raise EvidenceParseError("CSV evidence exceeds the safe column limit.")
            canonical = " | ".join(cell.strip() for cell in row)
            if not canonical.strip(" |"):
                continue
            segments.append(
                ParsedSegment(
                    ordinal=len(segments),
                    content=canonical,
                    start_offset=0,
                    end_offset=len(canonical),
                    locator={"kind": "csv_row", "row": row_number},
                    parser_name="csv-row-parser",
                )
            )
    except csv.Error as exc:
        raise EvidenceParseError("CSV evidence is malformed.") from exc
    return segments


def _parse_email(content: bytes) -> list[ParsedSegment]:
    try:
        message = BytesParser(policy=policy.default).parsebytes(content)
    except Exception as exc:  # noqa: BLE001 - email parser errors vary by malformed input
        raise EvidenceParseError("Email evidence is malformed.") from exc
    segments: list[ParsedSegment] = []
    headers = _email_headers(message)
    if headers:
        segments.append(
            ParsedSegment(
                ordinal=0,
                content=headers,
                start_offset=0,
                end_offset=len(headers),
                locator={"kind": "email_headers"},
                parser_name="rfc822-email-parser",
            )
        )
    body_index = 0
    parts = message.walk() if message.is_multipart() else [message]
    for part in parts:
        if part.is_multipart() or part.get_content_disposition() == "attachment":
            continue
        media_type = part.get_content_type()
        if media_type not in {"text/plain", "text/html"}:
            continue
        try:
            body = part.get_content()
        except (LookupError, UnicodeDecodeError) as exc:
            raise EvidenceParseError("Email body encoding is unsupported.") from exc
        body_text = str(body) if media_type == "text/plain" else _html_to_text(str(body))
        body_text = body_text.strip()
        if not body_text:
            continue
        segments.append(
            ParsedSegment(
                ordinal=len(segments),
                content=body_text,
                start_offset=0,
                end_offset=len(body_text),
                locator={"kind": "email_body", "part": body_index, "media_type": media_type},
                parser_name="rfc822-email-parser",
            )
        )
        body_index += 1
    return segments


def _email_headers(message: Message) -> str:
    labels = (
        ("From", "From"),
        ("To", "To"),
        ("Cc", "Cc"),
        ("Subject", "Subject"),
        ("Date", "Date"),
    )
    return "\n".join(
        f"{label}: {str(message.get(header)).strip()}"
        for header, label in labels
        if message.get(header)
    )


class _TextOnlyHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        clean = " ".join(data.split())
        if clean:
            self.parts.append(clean)


def _html_to_text(value: str) -> str:
    parser = _TextOnlyHTMLParser()
    parser.feed(value)
    parser.close()
    return "\n".join(parser.parts)


def _parse_pdf(content: bytes) -> list[ParsedSegment]:
    try:
        reader = PdfReader(io.BytesIO(content), strict=True)
    except PdfReadError as exc:
        raise EvidenceParseError("PDF evidence is malformed.") from exc
    if reader.is_encrypted:
        raise EvidenceParseError("Encrypted PDF evidence is not supported.")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise EvidenceParseError("PDF evidence exceeds the safe page limit.")
    segments: list[ParsedSegment] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        segments.append(
            ParsedSegment(
                ordinal=len(segments),
                content=text,
                start_offset=0,
                end_offset=len(text),
                locator={"kind": "pdf_page", "page": page_number},
                parser_name="pypdf-page-parser",
            )
        )
    return segments


def _parse_docx(content: bytes) -> list[ParsedSegment]:
    _validate_zip_container(content)
    try:
        document = Document(io.BytesIO(content))
    except (ValueError, KeyError, zipfile.BadZipFile) as exc:
        raise EvidenceParseError("DOCX evidence is malformed.") from exc
    segments: list[ParsedSegment] = []
    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            segments.append(
                ParsedSegment(
                    ordinal=len(segments),
                    content=text,
                    start_offset=0,
                    end_offset=len(text),
                    locator={"kind": "docx_paragraph", "paragraph": paragraph_number},
                    parser_name="docx-structure-parser",
                )
            )
    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            text = " | ".join(cell.text.strip() for cell in row.cells)
            if text.strip(" |"):
                segments.append(
                    ParsedSegment(
                        ordinal=len(segments),
                        content=text,
                        start_offset=0,
                        end_offset=len(text),
                        locator={
                            "kind": "docx_table_row",
                            "table": table_number,
                            "row": row_number,
                        },
                        parser_name="docx-structure-parser",
                    )
                )
    return segments


def _validate_zip_container(content: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ARCHIVE_FILES:
                raise EvidenceParseError("Document archive exceeds the safe file limit.")
            expanded_size = sum(entry.file_size for entry in entries)
            if expanded_size > MAX_ARCHIVE_EXPANDED_BYTES:
                raise EvidenceParseError("Document archive exceeds the safe expansion limit.")
            if any(".." in Path(entry.filename).parts for entry in entries):
                raise EvidenceParseError("Document archive contains an unsafe path.")
    except zipfile.BadZipFile as exc:
        raise EvidenceParseError("DOCX evidence is malformed.") from exc


def _parse_image(content: bytes, suffix: str, file_name: str) -> list[ParsedSegment]:
    expected_format = "JPEG" if suffix in {"jpg", "jpeg"} else "PNG"
    try:
        with Image.open(io.BytesIO(content)) as image:
            image.verify()
        with Image.open(io.BytesIO(content)) as image:
            width, height = image.size
            actual_format = image.format
    except (UnidentifiedImageError, OSError, Image.DecompressionBombError) as exc:
        raise EvidenceParseError("Image evidence is malformed or unsafe.") from exc
    if actual_format != expected_format:
        raise EvidenceParseError("Image content does not match its file extension.")
    if width * height > MAX_IMAGE_PIXELS:
        raise EvidenceParseError("Image evidence exceeds the safe pixel limit.")
    descriptor = f"Visual evidence: {Path(file_name).name} ({width}×{height} {expected_format})"
    return [
        ParsedSegment(
            ordinal=0,
            content=descriptor,
            start_offset=0,
            end_offset=len(descriptor),
            locator={"kind": "whole_image", "width": width, "height": height},
            parser_name="verified-image-parser",
            modality="image",
        )
    ]
